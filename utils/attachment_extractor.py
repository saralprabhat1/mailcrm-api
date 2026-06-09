# utils/attachment_extractor.py
# PURPOSE: Extract plain text from PDF and Word (.docx) email attachments.
# Phase 15 — called before AI extraction when hasAttachments is True.
#
# WHY THIS EXISTS:
#   Many requirement emails say "please find attached" with a blank or vague body.
#   The actual data (client, role, rates, duration, deadline) is in the PDF or Word
#   file. By extracting that text and appending it to the email body, we let the
#   AI extractor see the full content instead of returning all NULLs.
#
# EXTRACTION STRATEGY:
#   PDF  → try pdfplumber first (best for clean digital PDFs with text layer)
#        → fallback to PyMuPDF (fitz) if pdfplumber returns empty or errors
#          (handles more PDF formats; good for complex layouts)
#   DOCX → python-docx (paragraph text + table cell text — proposals often put
#          the actual role/rate/headcount data in tables, not paragraphs)
#   XLSX → openpyxl (reads every sheet, every row, all cell values as plain text)
#   DOC  → not supported (legacy format; would need LibreOffice) — skipped with warning
#   Other→ skipped silently
#
# SAFETY RULE: NEVER crash the pipeline.
#   Any exception → log a warning → return "" → caller continues normally.

import re
import logging

logger = logging.getLogger(__name__)

# Maximum PDF pages to read — prevents very long documents from flooding the AI prompt
MAX_PDF_PAGES = 20


def extract_text_from_attachment(attachment_name: str, attachment_bytes: bytes) -> str:
    """
    Extract plain text from a PDF or Word (.docx) email attachment.

    Parameters:
        attachment_name  — original file name, e.g. "RFQ_Drilling_Supervisor.pdf"
        attachment_bytes — raw bytes of the file (already downloaded from Graph API)

    Returns:
        Extracted plain text string, or "" on failure / unsupported type.
        Never raises an exception.
    """
    if not attachment_bytes:
        return ""

    name_lower = attachment_name.lower().strip()

    if name_lower.endswith(".pdf"):
        return _extract_pdf(attachment_name, attachment_bytes)

    elif name_lower.endswith(".docx"):
        return _extract_docx(attachment_name, attachment_bytes)

    elif name_lower.endswith(".xlsx"):
        return _extract_xlsx(attachment_name, attachment_bytes)

    elif name_lower.endswith(".zip"):
        return _extract_zip(attachment_name, attachment_bytes)

    elif name_lower.endswith(".doc"):
        # Legacy .doc binary format — python-docx cannot open these.
        # Would require LibreOffice or antiword for conversion, which are not
        # available in this environment.
        logger.warning(
            f"Attachment '{attachment_name}' is a legacy .doc file — "
            "text extraction not supported. Skipping."
        )
        print(f"         [attachment] '{attachment_name}' is .doc (legacy) — extraction not supported, skipping.")
        return ""

    else:
        # Everything else (.xlsx, .pptx, .txt, .zip, etc.) — skip silently.
        # These would need format-specific parsers we have not implemented.
        logger.debug(f"Attachment '{attachment_name}' type not supported for extraction. Skipping.")
        return ""


# ---------------------------------------------------------------------------
# PDF extraction
# ---------------------------------------------------------------------------

def _extract_pdf(name: str, data: bytes) -> str:
    """
    Extract text from a PDF.

    Tries pdfplumber first because it handles text-layer (searchable/digital)
    PDFs well and produces cleaner output. Falls back to PyMuPDF (fitz) if
    pdfplumber returns nothing — fitz handles a wider range of PDF variants.
    """
    # --- Attempt 1: pdfplumber ---
    try:
        import pdfplumber
        import io

        with pdfplumber.open(io.BytesIO(data)) as pdf:
            pages_text = []
            # Cap the number of pages so huge documents don't blow up the AI prompt
            for page in pdf.pages[:MAX_PDF_PAGES]:
                page_text = page.extract_text()
                if page_text:
                    pages_text.append(page_text)

        text = "\n".join(pages_text)

        if text.strip():
            print(f"         [attachment] pdfplumber: extracted {len(text)} chars from '{name}'")
            return _clean_whitespace(text)

        # pdfplumber returned empty — try the fallback
        print(f"         [attachment] pdfplumber returned empty for '{name}' — trying PyMuPDF...")

    except Exception as e:
        print(f"         [attachment] pdfplumber error on '{name}': {e} — trying PyMuPDF...")

    # --- Attempt 2: PyMuPDF (fitz) fallback ---
    return _extract_pdf_fitz(name, data)


def _extract_pdf_fitz(name: str, data: bytes) -> str:
    """PyMuPDF fallback for PDFs that pdfplumber cannot handle."""
    try:
        import fitz  # PyMuPDF
        import io

        doc = fitz.open(stream=data, filetype="pdf")
        pages_text = []
        for page_num, page in enumerate(doc):
            if page_num >= MAX_PDF_PAGES:
                break
            pages_text.append(page.get_text())
        doc.close()

        text = "\n".join(pages_text)

        if text.strip():
            print(f"         [attachment] PyMuPDF: extracted {len(text)} chars from '{name}'")
            return _clean_whitespace(text)

        print(
            f"         [attachment] PyMuPDF also returned empty for '{name}'. "
            "Likely a scanned image PDF — text extraction not possible without OCR."
        )
        return ""

    except Exception as e:
        logger.warning(f"PyMuPDF failed on '{name}': {e}")
        print(f"         [attachment] PyMuPDF error on '{name}': {e}")
        return ""


# ---------------------------------------------------------------------------
# DOCX extraction
# ---------------------------------------------------------------------------

def _extract_docx(name: str, data: bytes) -> str:
    """
    Extract text from a Word .docx file.

    Reads all paragraphs, then all table cell text (each row joined with " | ").
    Proposal documents commonly put the actual role/rate/headcount data in
    tables rather than paragraphs, so both are needed for the AI to see it.
    """
    try:
        from docx import Document
        import io

        doc = Document(io.BytesIO(data))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

        table_lines = []
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    table_lines.append(" | ".join(cells))

        text = "\n".join(paragraphs + table_lines)

        if text.strip():
            print(f"         [attachment] python-docx: extracted {len(text)} chars from '{name}'")
            return _clean_whitespace(text)

        print(f"         [attachment] python-docx returned empty for '{name}'.")
        return ""

    except Exception as e:
        logger.warning(f"python-docx failed on '{name}': {e}")
        print(f"         [attachment] python-docx error on '{name}': {e}")
        return ""


# ---------------------------------------------------------------------------
# XLSX extraction
# ---------------------------------------------------------------------------

def _extract_xlsx(name: str, data: bytes) -> str:
    """
    Extract text from an Excel .xlsx file.

    Reads every sheet and every row, joining cell values with " | ".
    Each sheet is labelled so the AI knows which section the data came from.
    """
    try:
        import openpyxl
        import io

        wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True, read_only=True)

        sheet_lines = []
        for ws in wb.worksheets:
            sheet_lines.append(f"[Sheet: {ws.title}]")
            for row in ws.iter_rows(values_only=True):
                cells = [str(v).strip() for v in row if v is not None and str(v).strip()]
                if cells:
                    sheet_lines.append(" | ".join(cells))

        text = "\n".join(sheet_lines)

        if text.strip():
            print(f"         [attachment] openpyxl: extracted {len(text)} chars from '{name}'")
            return _clean_whitespace(text)

        print(f"         [attachment] openpyxl returned empty for '{name}'.")
        return ""

    except Exception as e:
        logger.warning(f"openpyxl failed on '{name}': {e}")
        print(f"         [attachment] openpyxl error on '{name}': {e}")
        return ""


# ---------------------------------------------------------------------------
# ZIP extraction
# ---------------------------------------------------------------------------

_ZIP_BOILERPLATE_KEYWORDS = [
    "gcoc", "general condition", "terms and condition", "terms & condition",
    "contract condition", "omanization", "hse policy", "insurance",
    "nda", "non-disclosure", "appendix", "annex",
]

_ZIP_PRIORITY_KEYWORDS = [
    "boq", "bill of quantity", "rate schedule", "rate card", "scope of work",
    "rfq", "requirement", "manpower", "personnel", "position",
]

# Phase 16 — increased from 3000; attachment-enriched emails need headroom for BOQ/scope tables
_ZIP_PER_FILE_MAX_CHARS = 8000


def _zip_file_priority(inner_name: str) -> int:
    """
    Return a sort key for ZIP inner files — lower is higher priority.

    RFQ/BOQ content files: 0  (process first — most likely to have rates/headcount)
    Standard docx/xlsx:    1  (process next)
    Legal boilerplate:     2  (process last — GCOC, Omanization, NDA, Appendix)
    """
    n = inner_name.lower()
    if any(kw in n for kw in _ZIP_PRIORITY_KEYWORDS):
        return 0
    if any(kw in n for kw in _ZIP_BOILERPLATE_KEYWORDS):
        return 2
    return 1


def _extract_zip(name: str, data: bytes) -> str:
    """
    Extract text from supported files inside a ZIP archive, entirely in memory.

    Files are sorted by relevance: RFQ/BOQ content first, legal boilerplate last.
    Each inner file is capped at _ZIP_PER_FILE_MAX_CHARS so one large legal
    document cannot crowd out the actual requirement data.
    """
    import zipfile
    import io

    try:
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            inner_names = zf.namelist()
            print(f"         [attachment] ZIP '{name}' contains {len(inner_names)} file(s): "
                  f"{', '.join(inner_names[:5])}{'...' if len(inner_names) > 5 else ''}")

            # Sort: RFQ/BOQ files first, generic next, legal boilerplate last
            sorted_names = sorted(inner_names, key=_zip_file_priority)

            sections = []
            for inner_name in sorted_names:
                inner_lower = inner_name.lower()

                # Only process supported formats; skip directories and other types
                if not any(inner_lower.endswith(ext) for ext in (".pdf", ".docx", ".xlsx")):
                    logger.debug(f"ZIP entry '{inner_name}' type not supported — skipping")
                    continue

                try:
                    inner_bytes = zf.read(inner_name)
                except Exception as e:
                    print(f"         [attachment] Could not read '{inner_name}' from ZIP: {e}")
                    continue

                # Use the base filename (strip any subdirectory path inside the ZIP)
                base_name = inner_name.split("/")[-1]

                if inner_lower.endswith(".pdf"):
                    text = _extract_pdf(base_name, inner_bytes)
                elif inner_lower.endswith(".docx"):
                    text = _extract_docx(base_name, inner_bytes)
                elif inner_lower.endswith(".xlsx"):
                    text = _extract_xlsx(base_name, inner_bytes)
                else:
                    text = ""

                if text:
                    # Cap each file so one large doc cannot crowd out the rest
                    if len(text) > _ZIP_PER_FILE_MAX_CHARS:
                        text = text[:_ZIP_PER_FILE_MAX_CHARS] + "\n[... file truncated ...]"
                    sections.append(f"[ZIP: {base_name}]\n{text}")

            if not sections:
                print(f"         [attachment] ZIP '{name}' — no extractable text found in any entry.")
                return ""

            combined = "\n\n".join(sections)
            print(f"         [attachment] ZIP '{name}' — extracted {len(combined)} chars from "
                  f"{len(sections)} file(s)")
            return _clean_whitespace(combined)

    except zipfile.BadZipFile:
        print(f"         [attachment] '{name}' is not a valid ZIP file — skipping.")
        return ""
    except Exception as e:
        logger.warning(f"ZIP extraction failed on '{name}': {e}")
        print(f"         [attachment] ZIP error on '{name}': {e}")
        return ""


# ---------------------------------------------------------------------------
# Shared whitespace cleanup
# ---------------------------------------------------------------------------

def _clean_whitespace(text: str) -> str:
    """
    Normalise extracted text so it is clean to read and does not inflate
    the AI prompt with blank lines.

    Steps:
      1. Strip trailing spaces from every line
      2. Collapse 3 or more consecutive blank lines into 2
      3. Strip overall leading/trailing whitespace
    """
    lines = [line.rstrip() for line in text.splitlines()]
    text = "\n".join(lines)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
